import os
import json
import io
from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

# ==========================================
# 目录与环境初始化
# ==========================================
os.makedirs("data/raw", exist_ok=True)       # 原始材料层
os.makedirs("data/compiled", exist_ok=True)  # 编译层数据
os.makedirs("data/routing", exist_ok=True)   # 路由层规则指令
os.makedirs("static", exist_ok=True)         # 前端页面

# 初始化默认的路由指令文件 (技能规则)
routing_file = "data/routing/INDEX.md"
if not os.path.exists(routing_file):
    with open(routing_file, "w", encoding="utf-8") as f:
        f.write("""# 知识库路由与问答技能指令 (INDEX)

## 1. 角色定义
你是一个智能知识库问答助手。你的任务是根据提供的【背景知识】准确回答用户的查询。

## 2. 行为准则
- **严格忠于事实与原貌**：你只能基于【背景知识】中提供的信息进行回答。不要自己提取数据或过度总结，直接把知识库原文中的内容原样返回给用户。
- **动态意图澄清**：如果用户仅仅输入了某个产品名称（例如“5G消息”或“5G新通话”），而没有明确具体问题，你需要根据【背景知识】中包含的该产品的相关主题进行反问。例如：“关于5G消息，我可以为您介绍它的【产品定位】、【资费情况】、【应用场景】或【开通方式】等，请问您想了解哪一方面？”（具体反问选项请根据实际背景知识内容动态生成）。
- **多轮对话能力**：请结合用户上下文历史中的代词（如“它”、“这个”、“资费呢”）来理解当前问题。
- **兜底策略**：如果用户询问与产品无关的闲聊问题，或者背景知识中没有相关答案，请用礼貌的语气明确告知用户：“您好，该问题不在我的知识库内容范围内，我暂时无法为您解答呢。”
""")

# 初始化 OpenAI 客户端 (适配阿里云百炼 DashScope 兼容接口)
ALIYUN_API_KEY = "sk-4cbf9e896bb848f6a5db7e132d1e5fe1"
# ALIYUN_API_KEY = "2ac281eefe3a4d2e8741a783ffdc3022.6uUjBIWlos9MvPVu"

client = OpenAI(
    api_key=ALIYUN_API_KEY, 
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
    # base_url="https://open.bigmodel.cn/api/paas/v4"

)

app = FastAPI(title="五层架构知识库智能体", description="提供上传、编译、路由、输出、回写的全链路接口")

# ==========================================
# CORS 跨域配置 (支持客户前端调用)
# ==========================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有域名访问，生产环境建议替换为客户的实际域名，如 ["https://client.com"]
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有 HTTP 方法 (GET, POST 等)
    allow_headers=["*"],  # 允许所有请求头
)

# 挂载静态文件
app.mount("/static", StaticFiles(directory="static"), name="static")

# ==========================================
# 数据模型定义
# ==========================================
class UploadRequest(BaseModel):
    filename: str
    content: str

class CompileRequest(BaseModel):
    # 改为可选传参，如果不传或者传 "all"，则编译整个文件夹
    filename: str = "all"

class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    query: str
    history: list[Message] = []

# ==========================================
# 前端主页路由
# ==========================================
@app.get("/")
def read_index():
    return FileResponse("static/index.html")

# ==========================================
# 【1. 原始材料层】 接口：接收原始资料
# ==========================================
@app.post("/api/upload")
def upload_raw_material(req: UploadRequest):
    """将笔记、网页文本等保存为原始资料文件"""
    filepath = f"data/raw/{req.filename}.txt"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(req.content)
    return {"status": "success", "message": f"原始资料 {req.filename} 已保存"}

@app.post("/api/upload_file")
async def upload_document(files: list[UploadFile] = File(...)):
    """
    【1. 原始材料层】接收上传的单个或多个文档（TXT/PDF/DOCX），解析为纯文本并保存到原始资料层。
    实现“万物皆可入库”的思想。
    """
    saved_names = []
    
    for file in files:
        filename = file.filename
        ext = os.path.splitext(filename)[-1].lower()
        content = ""
        
        content_bytes = await file.read()
        
        try:
            if ext in [".txt", ".md", ".csv"]:
                content = content_bytes.decode("utf-8", errors="ignore")
            elif ext == ".pdf":
                if not PyPDF2:
                    continue # 略过未支持文件
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            elif ext == ".docx":
                if not docx:
                    continue
                doc = docx.Document(io.BytesIO(content_bytes))
                # 提取普通段落文本
                paras = [para.text for para in doc.paragraphs if para.text.strip()]
                # 提取表格文本
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            # 直接用空格分隔，不带任何格式，保持原文内容
                            tables_text.append(" ".join(row_data))
                
                content = "\n".join(paras) + "\n\n" + "\n".join(tables_text)
            else:
                continue # 不支持的格式直接跳过
        except Exception:
            continue # 解析失败跳过
            
        base_name = os.path.splitext(filename)[0]
        filepath = f"data/raw/{base_name}.txt"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        saved_names.append(base_name)
            
    return {"status": "success", "message": f"成功解析并存入 {len(saved_names)} 个文档", "saved_names": saved_names}

@app.get("/api/raw_files")
def list_raw_files():
    files = os.listdir("data/raw")
    return {"files": files}

# ==========================================
# 【2. 编译层】 接口：LLM处理原始材料
# ==========================================
@app.post("/api/compile")
def compile_material(req: CompileRequest):
    """
    【2. 编译层】LLM 对原始资料进行结构化拆解：
    支持单个文件或批量编译整个 raw 文件夹中的所有资料。
    """
    raw_dir = "data/raw"
    files_to_compile = []
    
    if req.filename == "all" or req.filename == "":
        # 编译所有文件
        if os.path.exists(raw_dir):
            files_to_compile = [f for f in os.listdir(raw_dir) if f.endswith(".txt")]
    else:
        # 编译指定文件
        if os.path.exists(f"{raw_dir}/{req.filename}"):
            files_to_compile = [req.filename]
            
    if not files_to_compile:
        raise HTTPException(status_code=404, detail="没有找到可编译的原始资料文件")

    compiled_results = []
    
    # 循环对每个文件进行编译 (在实际生产中，这里可以采用多线程/并发加速)
    for filename in files_to_compile:
        with open(f"{raw_dir}/{filename}", "r", encoding="utf-8") as f:
            raw_content = f.read()

        # 构建编译提示词
        prompt = f"""
        请你作为一个专业的知识库编译器。根据以下提供的【原始资料】，提取并生成结构化数据。
        请以 JSON 格式输出，包含以下字段：
        - summary: 该资料的核心摘要（50字以内）
        - entities: 提取出的核心实体列表（如人名、产品名、专有名词），用于建立关联图谱
        - index_keywords: 仔细阅读资料，提取出能代表资料各个段落核心主题的关键词列表（比如：定位、资费、案例、开通、功能等，尽量详细，不要遗漏）
        
        原始资料：
        {raw_content}
        
        注意：只输出合法的 JSON 字符串，不要包含 Markdown 格式的 ```json 标签。
        """
        
        try:
            response = client.chat.completions.create(
                # 严格使用你指定的阿里云模型名称
                model="qwen3-vl-plus",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1
            )

            compiled_result = response.choices[0].message.content.strip()
            # 清理可能携带的 markdown 标记
            if compiled_result.startswith("```json"):
                compiled_result = compiled_result[7:-3].strip()
                
            compiled_data = json.loads(compiled_result)
            compiled_data["raw_file"] = filename
            compiled_data["content"] = raw_content
            
            # 写入编译层
            compiled_path = f"data/compiled/{filename}.json"
            with open(compiled_path, "w", encoding="utf-8") as f:
                json.dump(compiled_data, f, ensure_ascii=False, indent=2)
                
            compiled_results.append(compiled_data)
        except Exception as e:
            print(f"文件 {filename} 编译失败: {str(e)}")
            continue

    return {
        "status": "success", 
        "message": f"成功编译了 {len(compiled_results)} 个文件", 
        "data": compiled_results
    }

# ==========================================
# 【3 & 4. 路由层 & 输出层】 接口：智能问答
# （此接口也是供外部调用的智能体问答接口）
# ==========================================
@app.post("/api/chat")
def chat_output(req: ChatRequest):
    """根据路由层指令寻找对应知识，并支持多轮对话输出"""
    query = req.query
    history = req.history
    
    # 路由层：读取外部的路由/技能指令文件
    # 每次强制重新写入 INDEX.md，以确保代码里最新的修改能实时生效
    with open("data/routing/INDEX.md", "w", encoding="utf-8") as f:
        f.write("""# 知识库路由与问答技能指令 (INDEX)

## 1. 角色定义
你是一个智能知识库问答助手。你的任务是根据提供的【背景知识】准确回答用户的查询。

## 2. 行为准则
- **严格忠于事实与原貌**：你只能基于【背景知识】中提供的信息进行回答。如果用户询问具体的条款、资费、案例等内容，请直接原样输出背景知识中的原文，不需要进行额外的数据提取、总结或重新组织。
- **动态意图澄清**：如果用户仅仅输入了某个产品名称（例如“5G消息”或“5G新通话”），而没有明确具体问题，你需要根据【背景知识】中包含的该产品的相关主题进行反问。例如：“关于5G消息，我可以为您介绍它的【产品定位】、【资费情况】、【应用场景】或【开通方式】等，请问您想了解哪一方面？”（具体反问选项请根据实际背景知识内容动态生成）。
- **多轮对话能力**：请结合用户上下文历史中的代词（如“它”、“这个”、“资费呢”）来理解当前问题。
- **兜底策略**：如果用户询问与产品无关的闲聊问题，或者背景知识中没有相关答案，请用礼貌的语气明确告知用户：“您好，该问题不在我的知识库内容范围内，我暂时无法为您解答呢。”
""")

    with open("data/routing/INDEX.md", "r", encoding="utf-8") as f:
        routing_instructions = f.read()
    
    # 检索逻辑：匹配编译层中的 index_keywords
    relevant_contexts = []
    if os.path.exists("data/compiled"):
        compiled_files = os.listdir("data/compiled")
        for c_file in compiled_files:
            with open(f"data/compiled/{c_file}", "r", encoding="utf-8") as f:
                data = json.load(f)
                
                # 提取历史记录和当前问题的所有文本作为匹配基准，提升上下文相关内容的命中率
                full_query_context = query + " " + " ".join([m.content for m in history])
                
                # 路由规则：如果用户的 query（或对话上下文） 包含了该文档的关键词，或者是实体
                keywords = data.get("index_keywords", []) + data.get("entities", [])
                
                # 匹配逻辑：文档关键词在用户的提问中，或者用户的提问在文档关键词中（双向匹配）
                is_match = False
                for kw in keywords:
                    if kw in full_query_context or full_query_context in kw:
                        is_match = True
                        break
                        
                # 如果是仅输入产品名称的澄清场景（字数少），也尽量带上包含该产品名的资料
                if len(query) <= 6:
                    for ent in data.get("entities", []):
                        if ent in query or query in ent:
                            is_match = True
                            break
                            
                if is_match:
                    relevant_contexts.append(f"【{data['raw_file']}】\n{data['content']}")
                
    context_str = "\n\n".join(relevant_contexts[:3]) if relevant_contexts else "未检索到相关知识。"
        
    # 组装输出层的 System Prompt (包含路由指令 + 检索到的知识)
    system_prompt = f"""{routing_instructions}
    
    【背景知识库内容】：
    {context_str}
    """
    
    # 组装多轮对话 Messages
    messages = [{"role": "system", "content": system_prompt}]
    
    # 将前端传来的历史对话放入模型上下文
    for msg in history:
        # 转换前端自定义的 role
        role = "assistant" if msg.role == "ai" else "user"
        # 过滤掉初始的寒暄语，防止污染上下文
        if "我已经准备好解答" not in msg.content:
            messages.append({"role": role, "content": msg.content})
            
    # 加入当前用户问题
    messages.append({"role": "user", "content": query})
    
    def generate():
        try:
            response = client.chat.completions.create(
                # 严格使用你指定的阿里云模型名称
                model="qwen3-vl-plus",
                messages=messages,
                temperature=0.3,
                stream=True # 开启大模型的流式输出
            )
            for chunk in response:
                # 兼容不同厂商的 chunk 格式，百炼的 stream 格式可能略有不同
                if hasattr(chunk, 'choices') and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    if hasattr(delta, 'content') and delta.content is not None:
                        yield delta.content
        except Exception as e:
            yield f"\n[问答失败: {str(e)}]"

    return StreamingResponse(generate(), media_type="text/event-stream")

if __name__ == "__main__":
    import uvicorn
    import nest_asyncio
    
    # 允许在 Jupyter Notebook 或其他已存在事件循环的环境中运行 uvicorn
    nest_asyncio.apply()
    
    # 换一个端口 8080 防止被占用
    print("服务正在启动，请在浏览器访问: http://localhost:8080")
    uvicorn.run(app, host="0.0.0.0", port=8080)
